import 代理ID
import json
import re
from docx import Document
from io import BytesIO
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

def getSNIS(inputurl):
    代理ID.dailiip()
    tha_list = {}
    for i in range(1, 30):
        html = 代理ID.openurl(inputurl + str(i)).decode('utf-8')   ####http://thzthz.cc/forum.php?mod=forumdisplay&fid=204&sortid=2&sortid=2&page=4
        # if i == 5:
        #     with open('html/thzla1.html', 'wb') as thzla:
        #         thzla.write(html.encode('utf-8'))
        soup = BeautifulSoup(html, 'lxml')
        a = soup.find_all('div',class_="c cl")
        for each in a:
            a_message = each.find('a')
            BTname = a_message['title']
            # 获取完整地址
            fullhttp = 'http://thzthz.cc/' + a_message['href']
            # 获取图片（暂缺）
            tha_list[BTname] = fullhttp
    return tha_list
    # 将数据保存为一个json文件
    # jsondata = json.dumps(tha_list)
    # with open('html/thzla1.json', 'a') as f:
    #     f.write(jsondata)


def getBTaddress(thz_dict,file_name):
    for each in thz_dict:
        # print(thz_dict[each])
        代理ID.dailiip()
        try:
            html = 代理ID.openurl(thz_dict[each]).decode('utf-8')
            soup_2 = BeautifulSoup(html, 'lxml')
            first_op = soup_2.find('ignore_js_op')
            # h获取第二个<ignore_js_op>,得到正反封面图
            try:
                second_op = first_op.find_next_sibling('ignore_js_op')
                imgfile = second_op.img['zoomfile']
            except KeyError:
                imgfile = first_op.img['zoomfile']
            if re.match(r'/data/attachment/forum/+[^\s]*',imgfile):
                imgfile = 'http://thzthz.cc' + imgfile
            # 获取下载页面链接
            try:
                third_op = second_op.find_next_sibling('ignore_js_op')
                download_page = third_op.a['href']
                full_page = 'http://thzthz.cc/' + download_page
            except Exception as emptyerror:
                print(repr(emptyerror))
                full_page = 'http://thzthz.cc/'
            # #BT的下载链接
            if full_page != 'http://thzthz.cc/':
                downloadhtml = 代理ID.openurl(full_page).decode('utf-8')
                soup_3 = BeautifulSoup(downloadhtml, 'lxml')
                BTurl = soup_3.find('div',style="padding-left:10px;").a['href']
            else:
                BTurl = ''
            # 链接合并为字典
            thz_full = {}
            thz_full['introduceurl'] = thz_dict[each]
            thz_full['imgurl'] = imgfile
            thz_full['downloadurl'] = full_page
            thz_full['BTurl'] = BTurl
            thz_dict[each] = thz_full
            print(thz_dict[each])
        except Exception as refuseerror:
            print(repr(refuseerror))
    print('搜集结束')
    with open('helloworld_python//HTML//'+file_name+'.json', 'w') as thz_full_json:
        thzdata = json.dumps(thz_dict)
        thz_full_json.write(thzdata)
    return thz_dict


def thz_docx(thz_dict,inputword):
    document = Document()
    document.add_heading(inputword, level=1)

    for each in thz_dict:
        print(each)
        p = document.add_paragraph(each)
        p.bold = True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        代理ID.dailiip()
        try:
            imageData = 代理ID.openurl(thz_dict[each]['imgurl'])
            image_io = BytesIO()
            image_io.write(imageData)
            image_io.seek(0)
            document.add_picture(image_io, width=Inches(5))
            image_io.close()
            p = document.add_paragraph(thz_dict[each]['BTurl'])
        except Exception as refuseerror:
            print(repr(refuseerror))
    document.save('helloworld_python//HTML//'+ inputword +'.docx')


if __name__ == '__main__':
    inputurl = input('请输入网址')
    inputword = input('请输入系列名')
    listurl = getSNIS(inputurl)
    thz_dict = getBTaddress(listurl,inputword)
    thz_docx(thz_dict,inputword)
