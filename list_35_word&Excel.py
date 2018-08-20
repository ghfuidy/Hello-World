import openpyxl
import docx
import logging
logging.disable(logging.CRITICAL)
logging.basicConfig(level=logging.DEBUG, filename='storage/lis_35_log.md')

def rowDataSet(ws, row):
    rowSet = set()
    for cell in ws[row]:
        if cell.value not in [None, '']:
            ###单元格不为空，将是数据加入集合
            rowSet.add(cell.value)
    return rowSet   ###返回集合

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr 
    tbl.remove(tr)

###输出第一个表格
def exportDataToWordTable1(ws, excelRow, doc, analysis):
    docRows = 2
    docCols = 7
    table = doc.add_table(rows = docRows, cols = docCols, style = 'Table Grid')
    
    for row in range(docRows):
        for col in range(docCols):
            table.rows[row].cells[col].text = str(ws.cell(row=excelRow, column=col+1).value)
        excelRow += 1
        
    table.rows[0].cells[1].text = '票房'
    table.rows[0].cells[2].text = '天数'
    table.rows[0].cells[3].text = '票房/天数'
    table.rows[1].cells[3].text = '{:.3f}'.format(float(table.rows[1].cells[1].text)/float(table.rows[1].cells[2].text))
    table.rows[0].cells[0].text = ''

    ###analysis data
    if float(table.rows[1].cells[4].text) > 0.85 and float(table.rows[1].cells[5].text) > 0.85:
        analysis['direction'] = '南'
    else:
        analysis['direction'] = '北'
    
    if float(table.rows[1].cells[6].text)  < 0.2:
        analysis['feeling'] = '凉'
    else:
        analysis['feeling'] = '暖'
    
    excelRow +=1 ###跳过factor_loading此行
    vSet = set()
    
    while ws.cell(row=excelRow, column=2).value != 'alpha': 
        ###factor_loading 数据放入vSet集合
        v = ws.cell(row=excelRow, column=2).value
        vSet.add(abs(v))
        excelRow += 1

    if min(vSet) > 0.3:
        analysis['season'] = '夏'
        analysis['duration'] = '长'
    else:
        analysis['season'] = '冬'
        analysis['duration'] = '短'
    
    return excelRow, analysis
    doc.add_paragraph('\n')
###输出第二个表格
def exportDataToWordTable2(ws, excelRow, doc, analysis):
    docCols = 2

    table = doc.add_table(rows=1, cols=docCols, style='Table Grid')
    wordRow = 0
    
    ###excelRow start from 'alpha' to 'X'
    while ws.cell(row=excelRow, column=1).value != 'X':
        for docCol in range(docCols):
            table.rows[wordRow].cells[docCol].text = str(ws.cell(row=excelRow, column=docCol+1).value)
        wordRow += 1
        excelRow += 1
        table.add_row()
        rowObj = table.add_row()
        ###删掉最后一行
        remove_row(table, rowObj)

    analysis['alpha'] = float(table.rows[1].cells[1].text)

    if analysis['alpha'] > 0.8:
        analysis['climate'] = '炎热'
    else:
        analysis['climate'] = '严寒'
    
    doc.add_paragraph('\n')
    return excelRow, analysis
###输出第三个表格
def exportDataToWordTable3(ws, excelRow, doc):
    docCols = len(rowDataSet(ws, excelRow))
    logging.debug('third excel of docCols is {}'.format(docCols))
    ###piant a excel
    table = doc.add_table(rows=docCols, cols=docCols, style='Table Grid')
    for docRow in range(docCols):
        for docCol in range(docCols):
            table.rows[docRow].cells[docCol].text = str(ws.cell(row=excelRow, column=docCol+1).value)
        excelRow += 1
    doc.add_paragraph('\n')
    return excelRow

###文字输出
def textOutput(doc, analysis):
    outputStr = '''人皆苦{}， 我爱{}日{}。 \n熏风自{}来， 殿阁生微{}。 \n
    alpha = {}'''.format(analysis['climate'], analysis['season'], analysis['duration'],
    analysis['direction'], analysis['feeling'], analysis['alpha'])

    doc.add_paragraph(outputStr)
    doc.add_paragraph('\n')

    doc.add_paragraph('-*-'*33)
    doc.add_page_break()


###start paragram
wb = openpyxl.load_workbook('storage/data.xlsx')
ws = wb['Sheet 1']

doc = docx.Document()

dataAnalysis = {'direction': '',
'feeling': '',
'season': '',
'duration': '',
'climate': '',
'alpha': 0}

###初始化Excel指针
row = 1

while row < ws.max_row:
    if {'TLI', 'CFI', 'RMSEA'} < rowDataSet(ws, row):
        ###输出第一张表。并分析数据
        row, dataAnalysis = exportDataToWordTable1(ws, row, doc, dataAnalysis)

        row, dataAnalysis = exportDataToWordTable2(ws, row, doc, dataAnalysis)

        row = exportDataToWordTable3(ws, row, doc)

        textOutput(doc, dataAnalysis)

    row += 1
doc.save('storage/Analysis.docx')