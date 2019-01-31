import 代理ID
import json
import re
from docx import Document
from io import BytesIO
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

# def thz_docx(name,image,torrent):
#     print(name)
#     p = document.add_paragraph(name)
#     p.bold = True
#     p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     try:
#         imageData = 代理ID.openurl(image)
#         image_io = BytesIO()
#         image_io.write(imageData)
#         image_io.seek(0)
#         document.add_picture(image_io, width=Inches(5))
#         image_io.close()
#         p = document.add_paragraph(torrent)
#     except Exception as refuseerror:
#         print('thz_docx',repr(refuseerror))


def getallpage(pagenum):
    page_dict = {} #全部页面的字典
    for i in range(1,pagenum+1):
        print(i)
        代理ID.dailiip()
        url = 'http://thzvv.net/forum-181-'+ str(i) +'.html'
        try:
            html = 代理ID.openurl(url).decode('utf-8')
            soup = BeautifulSoup(html,'lxml')
            list_table = soup.find('table',id='threadlisttableid')
            page_list = list_table.find_all('tbody',id=re.compile("normalthread_\d+"))
            for each in page_list:
                # print(each)
                url_dict = {} #地址的url字典
                page_url = 'http://thzvv.net/' + each.find('a')['href']
                page_name = each.find('a',class_='s xst').string
                url_dict['url'] = page_url
                page_dict[page_name] = url_dict #将url字典嵌套给全部页面字典
        except Exception as e:
            print('getallpage',repr(e))
            pass
    
    return page_dict
            
def get_page_info(pagelist):
    for each in pagelist:
        # print(pagelist[each]['url'])
        pageurl = pagelist[each]['url']
        代理ID.dailiip()
        try:
            html = 代理ID.openurl(pageurl).decode('utf-8')
            soup = BeautifulSoup(html,'lxml')
            jiedian = soup.find('ignore_js_op')
            jiedian_2 = jiedian.find_next('ignore_js_op')
            jiedian = jiedian.parent
            pageimages = jiedian.find_all('img',limit=3)
            torrenturl = jiedian_2.find('a')['href']
            print(pageimages,torrenturl)
            ## 直接写入文件
            p = document.add_paragraph(each)
            p.bold = True
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for eachimg in pageimages:
                try:
                    imageData = 代理ID.openurl(eachimg['file'])
                    image_io = BytesIO()
                    image_io.write(imageData)
                    image_io.seek(0)
                    document.add_picture(image_io, width=Inches(5))
                    image_io.close()
                except Exception as refuseerror:
                    print('thz_docx',repr(refuseerror))
            p = document.add_paragraph('http://thzvv.net/' + torrenturl)
        except Exception as e:
            print(each,'is','get_page_info',repr(e))
 


if __name__ == '__main__':
    document = Document()
    getnumber = int(input('input your need number of pages'))
    getdata = getallpage(getnumber)
    get_page_info(getdata)
    print('文件写入完成')
    document.save('helloworld_python//HTML//everyday.docx')   