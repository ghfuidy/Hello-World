from docx import Document
from io import BytesIO
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import 代理ID

document = Document()

document.add_heading('IPZ系列', level=1)
with open('html/IPZ.json', 'r') as thzjson:
    jsondata = thzjson.read().encode('utf-8')
    thz_dict = json.loads(jsondata)
    for each in thz_dict:
        print(each)
        p = document.add_paragraph(each)
        p.bold = True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        代理ID.dailiip()
        imageData = 代理ID.openurl(thz_dict[each]['imgurl'])
        image_io = BytesIO()
        image_io.write(imageData)
        image_io.seek(0)
        document.add_picture(image_io, width=Inches(6))
        image_io.close()
        p = document.add_paragraph(thz_dict[each]['BTurl'])
        document.add_page_break()
document.save('html/IPZ.docx')
