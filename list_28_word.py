# import docx

# document = docx.Document()

# paragraph = document.add_paragraph('作者： 苏轼', 'Subtitle')
# prior_paragraph = paragraph.insert_paragraph_before('水调歌头', 'Title')
# document.add_paragraph(
#     '''
#     明月几时有，把酒问青天。 
#     不知天上宫阙，今夕是何年？ 
#     我欲乘风归去，又恐琼楼玉宇， 
#     高处不胜寒。 
#     起舞弄清影，何似在人间！ 

#     转朱阁，低绮户，照无眠。 
#     不应有恨，何事长向别时圆？ 
#     人有悲欢离合，月有阴晴圆缺， 
#     此事古难全。 
#     但愿人长久，千里共婵娟。
#     ''', 'Body Text 2')

# # #Adding a heading
# # document.add_heading('The real meaning of the universe')
# # document.add_heading('Heading 0', level = 0)
# # document.add_heading('Heading 1', level = 1)
# # document.add_heading('Heading 2', level = 2)

# # for p in document.paragraphs:
# #     print(len(p.text))
# #     print(p.style.name)

##提取整个文件的某一类格式
# # from docx.enum.style import WD_STYLE_TYPE
# # styles =document.styles
# # print("\n".join([s.name for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]))

# paragraph2 = document.add_paragraph('流金岁月')
# paragraph2.style = 'List Bullet 2'
# paragraph3 = document.add_paragraph('似水流年')
# paragraph3.style = 'List Number 2'

# #character style
# paragraph4 = document.add_paragraph()
# paragraph4.add_run('No matter what we breed, ')
# paragraph4.add_run('we still are made of greed. ').blod = True
# paragraph4.add_run('This is my kindom come').italic = True
# run = paragraph4.add_run('This is my kindom come')
# run.style = 'Intense Emphasis'

# from docx.shared import Pt
# run.font.name = 'Sana'
# run.font.size = Pt(20)

# #Adding a page bread
# document.add_page_break()
# table = document.add_table(rows = 9, cols =10, style = 'Table Grid')
# cell_1 = table.cell(1,2)
# cell_2 = table.cell(4,7)
# cell_1.merge(cell_2)
# print()

# document.save('test.docx')


##opne a exist word
import docx
from docx.shared import Cm

document = docx.Document('test.docx')
table = document.tables[0]

for row, obj_row in enumerate(table.rows):
    for col, cell in enumerate(obj_row.cells):
        cell.text = '{}, {}'.format(row, col)
table.columns[4].width = Cm(4)
table.style = 'Colorful Grid Accent 5'

row = table.rows[1]
row.cells[0].text = 'elephant'
row.cells[1].text = 'shark' 

row_count = len(table.rows)
col_count = len(table.columns)
print('row = {}, column = {}'.format(row_count, col_count))

# row = table.add_row()
document.add_page_break()
document.add_picture('helloworld_python/red.jpg',width = Cm(5))
document.save('test.docx')



