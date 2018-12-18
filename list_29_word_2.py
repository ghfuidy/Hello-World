import docx
from docx.shared import Cm, Pt, RGBColor

recordset = [{'Qty':3,'Name':'Fish','Desc':'Tom'},
            {'Qty':8,'Name':'Chess','Desc':'Jerry'},
            {'Qty':5,'Name':'Bacon','Desc':'Garfiled'}]

document = docx.Document()
document.add_heading('Docunment Title', 0)
p = document.add_paragraph('A plain paragraph having some')
p.add_run('blod').blod = True
p.add_run('and some')
p.add_run('italic').italic = True
document.add_heading('Heading, level 1')
document.add_paragraph('Intense  quote', style = 'Intense Quote')

document.add_paragraph('first item in unordered list', style = 'List Bullet')
document.add_paragraph('first item in ordered list', style = 'List Number')

document.add_picture('helloworld_python/red.jpg' ,width = Cm(10))
document.add_page_break

table = document.add_table(rows = 1, cols = 3, style = 'Table Grid')
#'_Rows' object has no attribute 'cells'
har_cells = table.rows[0].cells
har_cells[0].text = 'Qty'
har_cells[1].text = 'Name'
har_cells[2].text = 'Desc'

for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item['Qty'])
    row_cells[1].text = item['Name']
    row_cells[2].text = item['Desc']

document.add_page_break()
from PIL import Image, ImageDraw
from io import BytesIO
#在图像上输入文字
from PIL import ImageFont

p = document.add_paragraph()
r = p.add_run()
image_size = 20
for x in range(255):
    im = Image.new('RGB', (image_size, image_size), 'white')
    draw_obj = ImageDraw.Draw(im)
    draw_obj.ellipse((0,0,image_size-1,image_size-1), fill = 255 - x)
    colorNum = 255 - x
    print(colorNum)
    font1 = ImageFont.truetype(r'C:\Windows\Fonts\times.ttf', 10)
    draw_obj.text((5, 5), str(colorNum), font = font1)
    fake_buf_file = BytesIO()
    im.save(fake_buf_file, "png")
    r.add_picture(fake_buf_file)
    fake_buf_file.close()
   
#一句话的渐变
document.add_page_break()
p = document.add_paragraph()
text = '一个人的命运当然要靠自我奋斗，但是也要考虑到历史的进程。'
for i, ch in enumerate(text):
    run = p.add_run(ch)
    font = run.font
    font.name = 'Silom'
    font.size = Pt(16)
    font.color.rgb = RGBColor(i*10%200+55,i*20%200+55,i*30%200+55)

document.save('Icando.docx')



